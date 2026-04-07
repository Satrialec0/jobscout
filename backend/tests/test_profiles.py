import io
import pytest


def test_extract_docx_returns_text():
    """DOCX extraction returns the paragraph text from the document."""
    from docx import Document
    from app.services.resume_parser import _extract_docx

    doc = Document()
    doc.add_paragraph("Jane Doe - Software Engineer")
    doc.add_paragraph("Skills: Python, SQL, AWS")
    buf = io.BytesIO()
    doc.save(buf)

    result = _extract_docx(buf.getvalue())

    assert "Jane Doe" in result
    assert "Python" in result


def test_extract_docx_empty_doc_returns_empty_string():
    from docx import Document
    from app.services.resume_parser import _extract_docx

    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)

    result = _extract_docx(buf.getvalue())

    assert result == ""


def test_extract_unsupported_type_raises():
    """Uploading a plain text file raises an HTTP 400 error."""
    import asyncio
    from fastapi import HTTPException
    from app.services.resume_parser import extract_resume_text

    class FakeUploadFile:
        content_type = "text/plain"
        filename = "resume.txt"

        async def read(self):
            return b"plain text content"

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(extract_resume_text(FakeUploadFile()))

    assert exc_info.value.status_code == 400
    assert "Unsupported file type" in exc_info.value.detail
